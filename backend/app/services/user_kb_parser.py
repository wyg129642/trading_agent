"""File parsers and chunker for the personal knowledge base.

Responsible for turning an uploaded file (.pdf / .md / .txt / .docx / .csv /
.html / .json) into plain-text-or-markdown content, then splitting that into
sliding-window chunks suitable for MongoDB ``$text`` retrieval.

PDF parsing prefers ``opendataloader-pdf`` (Java-backed, produces clean
markdown preserving headings, tables and reading order). If the JVM process
fails for any reason — typically on scanned or malformed PDFs — the parser
transparently falls back to ``pypdf`` so uploads never silently fail.

All functions in this module are **pure / sync** so they can be offloaded to
``asyncio.to_thread`` in the service layer without reentering the event loop.
"""

from __future__ import annotations

import json
import logging
import re
import subprocess  # noqa: S404 - used only for pinned opendataloader-pdf invocation
import tempfile
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


# ── File-type registry ─────────────────────────────────────────


# Declared extensions and their content types. The upload endpoint rejects
# anything outside this set before a byte hits Mongo, and the parser picks a
# concrete strategy per extension below. Keep this list conservative — every
# new type is an attack surface.
SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "application/pdf",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".txt": "text/plain",
    ".text": "text/plain",
    ".docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
    ".csv": "text/csv",
    ".json": "application/json",
    ".html": "text/html",
    ".htm": "text/html",
    # Audio — processed out-of-process by the jumpbox Qwen3-ASR service (see
    # user_kb_asr_client.py). No entry in _PARSERS: the dispatch in
    # user_kb_service._do_parse checks is_audio() and routes around parse_file.
    ".mp3": "audio/mpeg",
    ".wav": "audio/wav",
    ".m4a": "audio/mp4",
    ".flac": "audio/flac",
    ".ogg": "audio/ogg",
    ".opus": "audio/opus",
    ".webm": "audio/webm",
    ".aac": "audio/aac",
}


# Subset of SUPPORTED_EXTENSIONS that goes through the ASR pipeline. Kept
# here (not in the service module) so the file-type registry stays in one
# place and the parser knows not to dispatch to a sync _PARSERS entry.
AUDIO_EXTENSIONS: frozenset[str] = frozenset({
    ".mp3", ".wav", ".m4a", ".flac", ".ogg", ".opus", ".webm", ".aac",
})


def is_supported(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTENSIONS


def is_audio(filename: str) -> bool:
    """Return True iff the filename's extension routes through the ASR path."""
    return Path(filename).suffix.lower() in AUDIO_EXTENSIONS


def content_type_for(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext, "application/octet-stream")


# ── Result types ───────────────────────────────────────────────


@dataclass
class ParseResult:
    """Output of a successful parse — plain-ish markdown plus diagnostics."""

    text: str
    parser: str       # "opendataloader-pdf" | "pypdf" | "markdown" | ...
    warnings: list[str]


class ParseError(Exception):
    """Raised when a file cannot be parsed by any available backend."""


# ── Per-format parsers ─────────────────────────────────────────


def _read_utf8(data: bytes) -> str:
    """Decode bytes as UTF-8 with tolerant error handling."""
    # Try UTF-8 first (most modern files). Fall through to common CN encodings
    # before giving up with replacement — losing a few glyphs is strictly better
    # than refusing the upload.
    for enc in ("utf-8", "utf-8-sig", "gb18030", "gbk", "big5", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _parse_pdf_with_opendataloader(pdf_path: Path, out_dir: Path) -> str:
    """Run opendataloader-pdf (JVM) on a PDF, return extracted markdown.

    Raises ParseError if the JVM call fails or produces no markdown. The caller
    is responsible for falling back to the pypdf path.
    """
    try:
        import opendataloader_pdf  # type: ignore
    except Exception as e:  # pragma: no cover - explicit opt-out path
        raise ParseError(f"opendataloader-pdf unavailable: {e}") from e

    try:
        opendataloader_pdf.convert(
            input_path=[str(pdf_path)],
            output_dir=str(out_dir),
            format="markdown",
            quiet=True,
        )
    except subprocess.CalledProcessError as e:  # pragma: no cover - JVM failure
        raise ParseError(
            f"opendataloader-pdf JVM exited non-zero: {e}"
        ) from e
    except Exception as e:
        raise ParseError(f"opendataloader-pdf failed: {e}") from e

    # The library writes <stem>.md next to the PDF under out_dir, mirroring
    # any subdirectory structure. Walk for the first produced .md file.
    md_files = sorted(out_dir.rglob("*.md"))
    if not md_files:
        raise ParseError("opendataloader-pdf produced no markdown output")

    text_chunks: list[str] = []
    for md in md_files:
        try:
            text_chunks.append(md.read_text(encoding="utf-8", errors="replace"))
        except OSError as e:
            logger.warning("Could not read opendataloader output %s: %s", md, e)
    merged = "\n\n".join(t for t in text_chunks if t.strip())
    if not merged.strip():
        raise ParseError("opendataloader-pdf produced empty markdown")
    return merged


def _parse_pdf_with_pypdf(data: bytes) -> str:
    """Fallback PDF parser using pypdf — pure Python, no JVM, no layout.

    Less accurate than opendataloader (no reading-order heuristics, no tables,
    weaker on CJK PDFs), but always available and robust to JVM hiccups.
    """
    try:
        from pypdf import PdfReader
    except Exception as e:  # pragma: no cover - explicit missing dep
        raise ParseError(f"pypdf unavailable: {e}") from e

    import io
    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:  # pragma: no cover - malformed page
            logger.warning("pypdf page %d extract failed: %s", i, e)
            pages.append("")
    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if not text.strip():
        raise ParseError("pypdf extracted no text")
    return text


def _parse_pdf(data: bytes) -> ParseResult:
    warnings: list[str] = []
    with tempfile.TemporaryDirectory(prefix="user_kb_pdf_") as tmp:
        tmp_path = Path(tmp)
        pdf_path = tmp_path / "in.pdf"
        out_dir = tmp_path / "out"
        out_dir.mkdir()
        pdf_path.write_bytes(data)
        try:
            text = _parse_pdf_with_opendataloader(pdf_path, out_dir)
            return ParseResult(
                text=text, parser="opendataloader-pdf", warnings=warnings,
            )
        except ParseError as e:
            warnings.append(f"opendataloader-pdf failed: {e}; falling back to pypdf")
            logger.info("opendataloader failed on upload — falling back to pypdf: %s", e)
    # Out-of-tempdir fallback: pypdf reads the bytes directly.
    text = _parse_pdf_with_pypdf(data)
    return ParseResult(text=text, parser="pypdf", warnings=warnings)


def _parse_docx(data: bytes) -> ParseResult:
    """Extract text from a .docx file. Requires python-docx.

    python-docx isn't a hard dependency; if absent we raise ParseError so the
    upload is rejected with a clear "install docx support" message. This keeps
    the core feature useful even on a bare install.
    """
    try:
        import docx  # type: ignore
    except Exception as e:
        raise ParseError(
            "docx support not installed — `pip install python-docx` to enable"
        ) from e
    import io
    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.rstrip())
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    if not text.strip():
        raise ParseError("docx contained no extractable text")
    return ParseResult(text=text, parser="python-docx", warnings=[])


def _parse_markdown(data: bytes) -> ParseResult:
    return ParseResult(text=_read_utf8(data), parser="markdown", warnings=[])


def _parse_text(data: bytes) -> ParseResult:
    return ParseResult(text=_read_utf8(data), parser="text", warnings=[])


def _parse_csv(data: bytes) -> ParseResult:
    """Treat CSV as a human-readable pipe-separated table for retrieval."""
    import csv
    import io
    text = _read_utf8(data)
    reader = csv.reader(io.StringIO(text))
    rows = []
    for r in reader:
        cleaned = [c.strip() for c in r]
        rows.append(" | ".join(cleaned))
    return ParseResult(text="\n".join(rows), parser="csv", warnings=[])


def _parse_json(data: bytes) -> ParseResult:
    text = _read_utf8(data)
    try:
        parsed = json.loads(text)
        pretty = json.dumps(parsed, ensure_ascii=False, indent=2)
        return ParseResult(text=pretty, parser="json", warnings=[])
    except json.JSONDecodeError:
        # Store as-is — still searchable.
        return ParseResult(
            text=text, parser="json", warnings=["malformed JSON, stored verbatim"],
        )


def _parse_html(data: bytes) -> ParseResult:
    """Strip HTML to readable text using BeautifulSoup (already in requirements)."""
    try:
        from bs4 import BeautifulSoup
    except Exception as e:
        raise ParseError(f"bs4 unavailable: {e}") from e
    html = _read_utf8(data)
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()
    text = soup.get_text("\n", strip=True)
    if not text.strip():
        raise ParseError("HTML contained no extractable text")
    return ParseResult(text=text, parser="html", warnings=[])


_PARSERS: dict[str, "_ParserFn"] = {
    ".pdf": _parse_pdf,
    ".md": _parse_markdown,
    ".markdown": _parse_markdown,
    ".txt": _parse_text,
    ".text": _parse_text,
    ".docx": _parse_docx,
    ".csv": _parse_csv,
    ".json": _parse_json,
    ".html": _parse_html,
    ".htm": _parse_html,
}


# Public hook: lets downstream plugins register a parser for a new file type
# without editing this module. Example (to add .pptx support):
#
#     from backend.app.services.user_kb_parser import (
#         register_parser, ParseResult,
#     )
#     def _pptx(data: bytes) -> ParseResult:
#         ...
#     register_parser(".pptx", _pptx,
#         content_type="application/vnd.openxmlformats-officedocument."
#                      "presentationml.presentation")
#
# ``register_parser`` is safe to call at import time; subsequent uploads
# immediately pick it up via :func:`parse_file`.
from typing import Callable, Optional  # noqa: E402 - intentional after decls

_ParserFn = Callable[[bytes], ParseResult]


def register_parser(
    extension: str,
    parser: _ParserFn,
    *,
    content_type: Optional[str] = None,
) -> None:
    """Register a parser for a new file extension.

    :param extension: File extension including the dot, e.g. ``".pptx"``.
        Case-insensitive; always stored lowercase.
    :param parser: Callable taking the raw bytes and returning a
        :class:`ParseResult`. Should raise :class:`ParseError` on failure.
    :param content_type: Optional MIME string served back to clients via
        Content-Type on download. Falls back to ``application/octet-stream``.
    """
    ext = extension.lower()
    if not ext.startswith("."):
        raise ValueError(f"extension must start with '.': {extension!r}")
    _PARSERS[ext] = parser
    if content_type:
        SUPPORTED_EXTENSIONS[ext] = content_type
    elif ext not in SUPPORTED_EXTENSIONS:
        SUPPORTED_EXTENSIONS[ext] = "application/octet-stream"


def parse_file(filename: str, data: bytes) -> ParseResult:
    """Dispatch parse by file extension. Raises ParseError on unsupported type."""
    ext = Path(filename).suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise ParseError(f"unsupported file type: {ext or '(no extension)'}")
    if not data:
        raise ParseError("file is empty")
    return parser(data)


# ── Chunker ────────────────────────────────────────────────────


_PARAGRAPH_RE = re.compile(r"\n{2,}")
_WHITESPACE_RE = re.compile(r"[ \t]+")


def _normalize(text: str) -> str:
    """Collapse runs of whitespace but preserve paragraph breaks."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = _PARAGRAPH_RE.split(text)
    cleaned = []
    for p in paragraphs:
        p = _WHITESPACE_RE.sub(" ", p).strip()
        if p:
            cleaned.append(p)
    return "\n\n".join(cleaned)


def chunk_text(
    text: str,
    *,
    chunk_size: int = 1000,
    overlap: int = 200,
) -> list[str]:
    """Split text into roughly ``chunk_size``-char chunks with ``overlap`` overlap.

    The algorithm is paragraph-aware where possible:
      1. Split into paragraphs on blank lines.
      2. Greedily pack paragraphs into chunks up to ``chunk_size``.
      3. When a single paragraph exceeds ``chunk_size``, fall back to a sliding
         window of ``chunk_size`` chars with ``overlap`` at the tail.

    Returns a list of trimmed, non-empty chunks. Returns ``[]`` for empty
    input rather than a list containing ``""`` so callers can treat empty
    parse results uniformly.
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size must be > 0")
    if overlap < 0:
        raise ValueError("overlap must be >= 0")
    if overlap >= chunk_size:
        # An overlap equal to the chunk size produces zero forward progress
        # and the loop below never terminates. Clamp early with a warning.
        logger.warning(
            "chunk overlap %d >= chunk_size %d; clamping to %d",
            overlap, chunk_size, chunk_size // 4,
        )
        overlap = max(0, chunk_size // 4)

    normalized = _normalize(text)
    if not normalized:
        return []

    paragraphs = normalized.split("\n\n")

    chunks: list[str] = []
    buf: list[str] = []
    buf_len = 0

    def _flush_buf() -> None:
        if buf:
            chunks.append("\n\n".join(buf).strip())

    for para in paragraphs:
        if len(para) > chunk_size:
            # Flush the pending paragraph pack first.
            _flush_buf()
            buf = []
            buf_len = 0
            # Sliding window over the long paragraph.
            step = max(1, chunk_size - overlap)
            i = 0
            while i < len(para):
                window = para[i : i + chunk_size].strip()
                if window:
                    chunks.append(window)
                if i + chunk_size >= len(para):
                    break
                i += step
            continue

        # +2 accounts for the "\n\n" separator we'll emit between paragraphs.
        projected = buf_len + len(para) + (2 if buf else 0)
        if projected > chunk_size and buf:
            _flush_buf()
            buf = [para]
            buf_len = len(para)
        else:
            buf.append(para)
            buf_len = projected if buf_len else len(para)

    _flush_buf()
    return [c for c in chunks if c]
